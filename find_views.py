import docx

def find_views(path):
    try:
        doc = docx.Document(path)
        print(f"Read {path}")
        keywords = ["待办审批事务", "班费支出明细", "学生个人学业总览", "评优申报详情", "v_"]
        
        # We search through the document paragraphs and tables
        # Let's just print paragraphs that mention the views
        matches = []
        for p in doc.paragraphs:
            for kw in keywords:
                if kw in p.text:
                    matches.append(p.text.strip())
                    break
        
        # Also check tables
        for idx, t in enumerate(doc.tables):
            table_text = ""
            for r in t.rows:
                for c in r.cells:
                    table_text += c.text + " "
            for kw in keywords:
                if kw in table_text:
                    matches.append(f"Table {idx} contains {kw}")
                    break
                    
        # Just print out a summary
        for m in matches:
            if len(m) > 0:
                print(m[:100])
                
    except Exception as e:
        print(e)

if __name__ == "__main__":
    find_views(r"C:\Users\13238\Desktop\shujuku\物理结构设计及系统实现——第九组.docx")
